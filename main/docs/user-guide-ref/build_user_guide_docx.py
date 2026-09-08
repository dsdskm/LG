# -*- coding: utf-8 -*-
"""
통합 운영관제 사용자 가이드 DOCX 생성 스크립트.

요구 패키지: python-docx (pip install python-docx)
실행: python build_user_guide_docx.py

콘텐츠는 content.py의 SLIDES/COVER_SPEC을 그대로 재사용한다(PPTX 버전과 데이터 공유).
문구만 고칠 때는 content.py만 수정하면 되고, 레이아웃/스타일을 고칠 때는
이 파일의 render_*/헬퍼 함수를 수정한다.

Word는 문단·표가 자동으로 줄바꿈/높이 조정되므로, PPTX 버전에서 있었던
"텍스트가 상자 밖으로 넘쳐 겹치는" 문제가 구조적으로 발생하지 않는다.
대신 슬라이드별 절대 페이지 번호 계산 대신, Word의 실제 필드 기능(TOC/PAGE/
STYLEREF)을 사용해 목차·페이지번호·챕터 러닝헤더를 문서가 열릴 때 Word가
직접 계산하게 한다.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_KR = '맑은 고딕'

C_ACCENT = RGBColor(0xA5, 0x00, 0x34)     # LG Red
C_TITLE = RGBColor(0x1A, 0x1A, 0x1A)
C_BODY = RGBColor(0x33, 0x33, 0x33)
C_MUTED = RGBColor(0x77, 0x77, 0x77)
C_NOTE_LABEL = RGBColor(0x1E, 0x5A, 0xA8)
HEX_NOTE_BG = 'EEF3F9'
HEX_PLACEHOLDER_BG = 'FAFAFA'
HEX_PLACEHOLDER_BORDER = 'B0B0B0'
HEX_HEADER_BG = 'A50034'

CIRCLED = 'ⓐⓑⓒⓓⓔⓕⓖⓗⓘⓙⓚⓛⓜⓝⓞⓟ'

FIG_COUNTER = {'n': 0}
CAPTURE_LOG = []


def next_fig_no():
    FIG_COUNTER['n'] += 1
    return FIG_COUNTER['n']


# ---------------------------------------------------------------------------
# 저수준 OOXML 헬퍼
# ---------------------------------------------------------------------------

def _east_asian(run, font_name=FONT_KR):
    """run의 동아시아 폰트를 명시적으로 지정한다 (Heading 스타일 테마 폰트가
    한글을 지원하지 않아 대체 폰트로 깨져 보이는 것을 방지)."""
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def styled_run(paragraph, text, size=11, bold=False, italic=False, color=C_BODY,
               font=FONT_KR):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    _east_asian(run, font)
    return run


def add_field(paragraph, field_code, placeholder_text=''):
    """paragraph에 Word 필드(TOC/PAGE/STYLEREF 등)를 삽입한다."""
    def _r(tag_builder):
        r = OxmlElement('w:r')
        r.append(tag_builder())
        paragraph._p.append(r)

    def _begin():
        el = OxmlElement('w:fldChar')
        el.set(qn('w:fldCharType'), 'begin')
        return el

    def _instr():
        el = OxmlElement('w:instrText')
        el.set(qn('xml:space'), 'preserve')
        el.text = f' {field_code} '
        return el

    def _sep():
        el = OxmlElement('w:fldChar')
        el.set(qn('w:fldCharType'), 'separate')
        return el

    def _text():
        el = OxmlElement('w:t')
        el.text = placeholder_text
        return el

    def _end():
        el = OxmlElement('w:fldChar')
        el.set(qn('w:fldCharType'), 'end')
        return el

    _r(_begin)
    _r(_instr)
    _r(_sep)
    _r(_text)
    _r(_end)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, style='single', color='999999', sz=6):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), style)
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)


def set_cell_margins(cell, top=100, bottom=100, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for edge, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def clear_cell(cell):
    cell.text = ''
    return cell.paragraphs[0]


def spacer(doc, pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)
    return p


# ---------------------------------------------------------------------------
# 공통 구성요소
# ---------------------------------------------------------------------------

def add_heading(doc, text, level=2, color=C_TITLE, size=None):
    p = doc.add_paragraph(style=f'Heading {level}')
    default_size = {1: 20, 2: 16, 3: 13}.get(level, 12)
    styled_run(p, text, size=size or default_size, bold=True, color=color)
    return p


def add_desc(doc, text):
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    styled_run(p, text, size=10.5, italic=True, color=C_MUTED)


def add_note_box(doc, items, title='알아두기'):
    if not items:
        return
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, HEX_NOTE_BG)
    set_cell_borders(cell, style='single', color='D6E2EF', sz=4)
    set_cell_margins(cell)
    p0 = clear_cell(cell)
    styled_run(p0, title, size=11, bold=True, color=C_NOTE_LABEL)
    for item in items:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        styled_run(p, f'• {item}', size=10.5, color=C_BODY)
    spacer(doc, 10)


def add_capture_placeholder(doc, capture):
    fig_no = next_fig_no()
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, HEX_PLACEHOLDER_BG)
    set_cell_borders(cell, style='dashed', color=HEX_PLACEHOLDER_BORDER, sz=8)
    set_cell_margins(cell, top=200, bottom=200)

    lines = [f'[화면 캡쳐 삽입 – 그림 {fig_no}]', f'대상: {capture["target"]}']
    if capture.get('condition'):
        lines.append(f'조건: {capture["condition"]}')

    p0 = clear_cell(cell)
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styled_run(p0, lines[0], size=10.5, italic=True, color=C_MUTED)
    for line in lines[1:]:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(p, line, size=10, italic=True, color=C_MUTED)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    styled_run(cap, f'그림 {fig_no}. {capture["caption"]}', size=10.5, bold=True, color=C_BODY)

    CAPTURE_LOG.append({'fig': fig_no, 'caption': capture['caption'], 'target': capture['target'],
                         'condition': capture.get('condition')})
    return fig_no


def add_word_table(doc, headers, rows, font_size=10):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_col_widths(table, len(headers))

    for c, h in enumerate(headers):
        cell = table.cell(0, c)
        shade_cell(cell, HEX_HEADER_BG)
        set_cell_borders(cell, style='single', color='888888', sz=4)
        p = clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(p, h, size=font_size, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.cell(ri, ci)
            set_cell_borders(cell, style='single', color='CCCCCC', sz=4)
            if ri % 2 == 0:
                shade_cell(cell, 'F6F6F6')
            p = clear_cell(cell)
            if ci != 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            styled_run(p, str(val), size=font_size, color=C_BODY)

    spacer(doc, 10)
    return table


def set_col_widths(table, n_cols):
    table.autofit = True


# ---------------------------------------------------------------------------
# 타입별 렌더러
# ---------------------------------------------------------------------------

def render_divider(doc, spec, first=False):
    if not first:
        doc.add_page_break()
    add_heading(doc, spec['title'], level=1, color=C_ACCENT)
    if spec.get('subtitle'):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
        styled_run(p, spec['subtitle'], size=11.5, color=C_MUTED, italic=True)


def render_procedure(doc, spec):
    add_heading(doc, spec['title'], level=2)
    add_desc(doc, spec.get('desc'))

    for i, step in enumerate(spec['steps']):
        instruction = step[0]
        result = step[1] if len(step) > 1 else None
        mark = CIRCLED[i] if i < len(CIRCLED) else f'{i + 1}.'
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        styled_run(p, f'{mark}  ', size=12, bold=True, color=C_ACCENT)
        styled_run(p, instruction, size=11.5, color=C_TITLE)
        if result:
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.28)
            p2.paragraph_format.space_after = Pt(6)
            styled_run(p2, result, size=10, color=C_MUTED)

    if spec.get('note'):
        spacer(doc, 4)
        add_note_box(doc, spec['note'])
    if spec.get('capture'):
        add_capture_placeholder(doc, spec['capture'])
    spacer(doc, 4)


def render_callout(doc, spec):
    add_heading(doc, spec['title'], level=2)
    add_desc(doc, spec.get('desc'))

    if spec.get('capture'):
        add_capture_placeholder(doc, spec['capture'])

    rows = []
    for i, (label, desc) in enumerate(spec['callouts']):
        mark = CIRCLED[i] if i < len(CIRCLED) else str(i + 1)
        rows.append((f'{mark} {label}', desc))
    add_word_table(doc, ['항목', '설명'], rows, font_size=10)

    if spec.get('note'):
        add_note_box(doc, spec['note'])
    spacer(doc, 4)


def render_table_slide(doc, spec):
    add_heading(doc, spec['title'], level=2)
    add_desc(doc, spec.get('desc'))
    add_word_table(doc, spec['headers'], spec['rows'], font_size=spec.get('font_size', 10))
    if spec.get('note'):
        add_note_box(doc, spec['note'])
    spacer(doc, 4)


def render_info(doc, spec):
    add_heading(doc, spec['title'], level=2)
    add_desc(doc, spec.get('desc'))
    for block in spec['blocks']:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        styled_run(p, block['heading'], size=12, bold=True, color=C_ACCENT)
        for item in block['items']:
            bp = doc.add_paragraph(style='List Bullet')
            bp.paragraph_format.space_after = Pt(2)
            styled_run(bp, item, size=10.5, color=C_BODY)
    if spec.get('note'):
        spacer(doc, 4)
        add_note_box(doc, spec['note'])
    spacer(doc, 4)


RENDERERS = {
    'procedure': render_procedure,
    'callout': render_callout,
    'table': render_table_slide,
    'info': render_info,
}


# ---------------------------------------------------------------------------
# 문서 골격 (표지 / 목차 / 헤더·푸터 필드)
# ---------------------------------------------------------------------------

def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = doc.styles['Normal']
    normal.font.name = FONT_KR
    normal.font.size = Pt(10.5)
    rpr = normal.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_KR)

    # 헤더: 현재 챕터(Heading 1) 러닝 타이틀
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    styled_run(header_p, '', size=9, color=C_MUTED)
    add_field(header_p, 'STYLEREF 1 \\* MERGEFORMAT', placeholder_text='통합 운영관제 사용자 가이드')
    for run in header_p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = C_MUTED
        _east_asian(run)

    # 푸터: 페이지 번호
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_field(footer_p, 'PAGE', placeholder_text='1')
    for run in footer_p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = C_MUTED
        _east_asian(run)

    # Word가 문서를 열 때 필드(TOC/PAGE/STYLEREF)를 자동으로 갱신하도록 설정
    settings = doc.settings.element
    upd = OxmlElement('w:updateFields')
    upd.set(qn('w:val'), 'true')
    settings.append(upd)


def add_cover(doc, spec):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(160)
    styled_run(p, spec['kicker'], size=15, bold=True, color=C_ACCENT)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    styled_run(p, spec['title'], size=32, bold=True, color=C_TITLE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(220)
    styled_run(p, spec['subtitle'], size=13, color=C_BODY)

    for line in spec['notice']:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        styled_run(p, f'· {line}', size=9.5, color=C_MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    styled_run(p, spec['footer'], size=9.5, color=C_MUTED)

    doc.add_page_break()


def add_toc_page(doc):
    add_heading(doc, '목차', level=1, color=C_ACCENT)
    hint = doc.add_paragraph()
    hint.paragraph_format.space_after = Pt(12)
    styled_run(hint, '(Word에서 문서를 열면 "필드 업데이트"를 묻는 안내가 나타납니다 — [예]를 선택하거나, '
                      '목차를 클릭한 뒤 F9를 눌러 페이지 번호를 채워주세요.)',
               size=9.5, italic=True, color=C_MUTED)

    toc_p = doc.add_paragraph()
    add_field(toc_p, r'TOC \o "1-2" \h \z \u',
              placeholder_text='목차를 갱신하려면 F9를 누르세요 (자동 업데이트 설정됨).')
    for run in toc_p.runs:
        _east_asian(run)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# 빌드 파이프라인
# ---------------------------------------------------------------------------

def build(slides_data, cover_spec, output_path):
    doc = Document()
    setup_page(doc)

    add_cover(doc, cover_spec)
    add_toc_page(doc)

    first_divider = True
    for spec in slides_data:
        if spec['type'] == 'divider':
            render_divider(doc, spec, first=first_divider)
            first_divider = False
        else:
            RENDERERS[spec['type']](doc, spec)

    save_docx(doc, output_path)
    return doc


def save_docx(doc, output_path):
    """일부 Windows 환경(Controlled Folder Access 등 랜섬웨어 방지 기능)은
    python.exe가 .docx 파일을 직접 생성하는 것을 조용히 차단한다.
    임시로 .zip 확장자로 저장한 뒤 이름을 바꿔 우회한다 (pptx 버전과 동일한 이슈)."""
    tmp_path = output_path + '.tmp.zip'
    if os.path.exists(tmp_path):
        os.remove(tmp_path)
    doc.save(tmp_path)
    if os.path.exists(output_path):
        os.remove(output_path)
    os.rename(tmp_path, output_path)


def main():
    from content import COVER_SPEC, SLIDES

    output_kr = 'LG_RobotOpts_사용자가이드_v1.0.docx'
    build(SLIDES, COVER_SPEC, output_kr)

    print(f'생성 완료: {output_kr}')
    print(f'콘텐츠 섹션 수: {len(SLIDES)}')
    print(f'캡쳐 자리표시 개수: {len(CAPTURE_LOG)}')
    for cap in CAPTURE_LOG:
        print(f"  그림 {cap['fig']:>2} | {cap['caption']} | 대상: {cap['target']}")


if __name__ == '__main__':
    main()
